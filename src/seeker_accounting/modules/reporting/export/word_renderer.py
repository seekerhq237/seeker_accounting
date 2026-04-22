"""Word (.docx) renderer for financial statements using python-docx.

Fully independent from ``seeker_accounting.platform.printing``.
Each financial statement type is rendered with intentional, statement-specific
formatting through the section-based layout system.
"""

from __future__ import annotations

import datetime
from decimal import Decimal
from typing import Sequence

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor
from docx.table import Table

from seeker_accounting.modules.reporting.export.export_models import (
    StatementCompanyInfo,
    StatementExportRow,
    StatementPageOrientation,
    StatementPageSize,
    StatementTableSection,
)

_BRAND_PRIMARY = RGBColor(0x1E, 0x3A, 0x5F)
_BRAND_LIGHT = RGBColor(0xF0, 0xF3, 0xF7)
_BRAND_BORDER = RGBColor(0xD0, 0xD7, 0xDE)
_SECTION_BG = "EDF1F6"
_SUBTOTAL_BG = "F5F6F8"
_TOTAL_BG = "E6EAF0"
_HIGHLIGHT_BG = "FFFCE8"
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_BLACK = RGBColor(0x1A, 0x1A, 0x1A)
_GRAY = RGBColor(0x88, 0x88, 0x88)
_ZERO = Decimal("0.00")


class FinancialStatementWordRenderer:
    """Renders financial statements to Word documents.

    Accepts ``StatementTableSection`` objects so that each statement
    can define its own column structure and visual layout.
    """

    def render(
        self,
        *,
        title: str,
        subtitle: str | None,
        date_label: str,
        company: StatementCompanyInfo,
        sections: Sequence[StatementTableSection],
        summary_pairs: Sequence[tuple[str, str]],
        output_path: str,
        page_size: StatementPageSize,
        orientation: StatementPageOrientation,
    ) -> None:
        doc = self._create_document(page_size, orientation)
        self._add_company_header(doc, company, page_size)
        self._add_title_block(doc, title, subtitle, date_label, page_size)

        for section in sections:
            self._add_section(doc, section, page_size)

        if summary_pairs:
            self._add_summary_block(doc, summary_pairs, page_size)
        self._add_footer(doc, page_size)
        doc.save(output_path)

    # ------------------------------------------------------------------
    # Document setup
    # ------------------------------------------------------------------

    def _create_document(
        self,
        page_size: StatementPageSize,
        orientation: StatementPageOrientation,
    ) -> Document:
        doc = Document()

        # Remove default empty paragraph
        if doc.paragraphs:
            p = doc.paragraphs[0]._element
            p.getparent().remove(p)

        section = doc.sections[0]
        w_mm, h_mm = page_size.dimensions_mm
        margin = page_size.margin_mm

        if orientation == StatementPageOrientation.LANDSCAPE:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = Mm(h_mm)
            section.page_height = Mm(w_mm)
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = Mm(w_mm)
            section.page_height = Mm(h_mm)

        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)

        return doc

    # ------------------------------------------------------------------
    # Company header
    # ------------------------------------------------------------------

    def _add_company_header(
        self,
        doc: Document,
        company: StatementCompanyInfo,
        page_size: StatementPageSize,
    ) -> None:
        font_pt = page_size.body_font_pt

        p = doc.add_paragraph()
        run = p.add_run(company.name)
        run.font.size = Pt(font_pt + 4)
        run.font.bold = True
        run.font.color.rgb = _BRAND_PRIMARY
        p.paragraph_format.space_after = Pt(0)

        if company.legal_name and company.legal_name != company.name:
            p2 = doc.add_paragraph()
            run2 = p2.add_run(company.legal_name)
            run2.font.size = Pt(font_pt)
            run2.font.color.rgb = _GRAY
            p2.paragraph_format.space_after = Pt(0)

        for line in company.address_block:
            pa = doc.add_paragraph()
            ra = pa.add_run(line)
            ra.font.size = Pt(font_pt - 1)
            ra.font.color.rgb = _GRAY
            pa.paragraph_format.space_after = Pt(0)

        id_parts: list[str] = []
        if company.tax_identifier:
            id_parts.append(f"Tax ID: {company.tax_identifier}")
        if company.registration_number:
            id_parts.append(f"Reg: {company.registration_number}")
        if id_parts:
            pi = doc.add_paragraph()
            ri = pi.add_run(" · ".join(id_parts))
            ri.font.size = Pt(font_pt - 1)
            ri.font.color.rgb = _GRAY
            pi.paragraph_format.space_after = Pt(4)

    # ------------------------------------------------------------------
    # Title block
    # ------------------------------------------------------------------

    def _add_title_block(
        self,
        doc: Document,
        title: str,
        subtitle: str | None,
        date_label: str,
        page_size: StatementPageSize,
    ) -> None:
        font_pt = page_size.body_font_pt

        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.size = Pt(font_pt + 3)
        run.font.bold = True
        run.font.color.rgb = _BRAND_PRIMARY
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)

        if subtitle:
            ps = doc.add_paragraph()
            rs = ps.add_run(subtitle)
            rs.font.size = Pt(font_pt)
            rs.font.color.rgb = _GRAY
            ps.paragraph_format.space_after = Pt(0)

        pd = doc.add_paragraph()
        rd = pd.add_run(date_label)
        rd.font.size = Pt(font_pt - 1)
        rd.font.color.rgb = _GRAY
        pd.paragraph_format.space_after = Pt(6)

        _add_bottom_border(pd, "1E3A5F", "12")

    # ------------------------------------------------------------------
    # Section (heading + table)
    # ------------------------------------------------------------------

    def _add_section(
        self,
        doc: Document,
        section: StatementTableSection,
        page_size: StatementPageSize,
    ) -> None:
        font_pt = page_size.body_font_pt

        if section.heading:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(section.heading)
            run.font.size = Pt(font_pt + 1)
            run.font.bold = True
            run.font.color.rgb = _BRAND_PRIMARY

        self._add_statement_table(doc, section.column_headers, section.rows, page_size)

    # ------------------------------------------------------------------
    # Statement table
    # ------------------------------------------------------------------

    def _add_statement_table(
        self,
        doc: Document,
        column_headers: Sequence[str],
        rows: Sequence[StatementExportRow],
        page_size: StatementPageSize,
    ) -> None:
        font_pt = page_size.body_font_pt
        col_count = len(column_headers)
        table = doc.add_table(rows=1, cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        hdr_row = table.rows[0]
        for i, hdr_text in enumerate(column_headers):
            cell = hdr_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(hdr_text)
            run.font.size = Pt(font_pt - 1)
            run.font.bold = True
            run.font.color.rgb = _WHITE
            if i >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_cell_shading(cell, "1E3A5F")

        # Data rows
        for row_data in rows:
            row_cells = table.add_row().cells
            kind = row_data.row_kind
            is_section = kind == "section"
            is_total = kind == "total"
            is_bold = kind in ("section", "group", "formula", "subtotal", "subsection", "total")
            is_formula = kind in ("formula", "subtotal")

            # Ref cell
            ref_text = row_data.ref_code if not is_section else ""
            self._set_cell_text(row_cells[0], ref_text, font_pt - 1, bold=False, color=_GRAY)

            # Label cell with indent
            indent_str = "    " * row_data.indent_level
            self._set_cell_text(
                row_cells[1],
                indent_str + row_data.label,
                font_pt,
                bold=is_bold,
            )

            # Amount cells
            for i in range(2, col_count):
                amt_idx = i - 2
                val = row_data.amounts[amt_idx] if amt_idx < len(row_data.amounts) else None
                self._set_cell_text(
                    row_cells[i],
                    self._fmt_amount(val),
                    font_pt,
                    bold=is_bold,
                    alignment=WD_ALIGN_PARAGRAPH.RIGHT,
                )

            # Row shading per kind
            if is_total:
                for cell in row_cells:
                    _set_cell_shading(cell, _TOTAL_BG)
            elif is_section:
                for cell in row_cells:
                    _set_cell_shading(cell, _SECTION_BG)
            elif is_formula:
                for cell in row_cells:
                    _set_cell_shading(cell, _SUBTOTAL_BG)
            elif row_data.is_highlight:
                for cell in row_cells:
                    _set_cell_shading(cell, _HIGHLIGHT_BG)

    # ------------------------------------------------------------------
    # Summary block
    # ------------------------------------------------------------------

    def _add_summary_block(
        self,
        doc: Document,
        pairs: Sequence[tuple[str, str]],
        page_size: StatementPageSize,
    ) -> None:
        font_pt = page_size.body_font_pt
        doc.add_paragraph()  # spacing

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for idx, (label, value) in enumerate(pairs):
            row = table.add_row()
            is_last = idx == len(pairs) - 1
            self._set_cell_text(row.cells[0], label, font_pt, bold=True)
            self._set_cell_text(
                row.cells[1],
                value,
                font_pt,
                bold=is_last,
                alignment=WD_ALIGN_PARAGRAPH.RIGHT,
            )
            if is_last:
                _set_cell_shading(row.cells[0], "F0F3F7")
                _set_cell_shading(row.cells[1], "F0F3F7")

    # ------------------------------------------------------------------
    # Footer
    # ------------------------------------------------------------------

    def _add_footer(self, doc: Document, page_size: StatementPageSize) -> None:
        font_pt = page_size.body_font_pt
        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Seeker Accounting · Generated {now}")
        run.font.size = Pt(font_pt - 2)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        p.paragraph_format.space_before = Pt(10)

    # ------------------------------------------------------------------
    # Cell helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _set_cell_text(
        cell,
        text: str,
        font_pt: int,
        *,
        bold: bool = False,
        color: RGBColor | None = None,
        alignment: int | None = None,
    ) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.size = Pt(font_pt)
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color
        if alignment is not None:
            p.alignment = alignment

    @staticmethod
    def _fmt_amount(value: Decimal | None) -> str:
        if value is None:
            return ""
        if value == _ZERO:
            return "–"
        if value < 0:
            return f"({abs(value):,.2f})"
        return f"{value:,.2f}"


# ---------------------------------------------------------------------------
# XML helpers for python-docx
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color_hex: str) -> None:
    """Set the background shading of a table cell."""
    tc = cell._element
    tc_pr = tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        from lxml import etree
        shading = etree.SubElement(tc_pr, qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def _add_bottom_border(paragraph, color_hex: str, size: str = "8") -> None:
    """Add a bottom border to a paragraph."""
    from lxml import etree

    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
