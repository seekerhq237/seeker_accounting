"""Payroll Export Service — orchestrates PDF payslip and CSV summary file exports.

This service owns:
- payslip PDF generation (single and batch)
- payroll summary CSV generation
- payroll summary PDF generation
- safe filename derivation
- audit event emission for export actions

All payroll truth comes from PayrollPrintService / PayrollRunService.
No new truth tables or stored rendered documents.
"""
from __future__ import annotations

import csv
import html
import io
import os
import re
from decimal import Decimal
from typing import Callable

from sqlalchemy.orm import Session

from seeker_accounting.db.unit_of_work import UnitOfWorkFactory
from seeker_accounting.modules.administration.services.permission_service import PermissionService
from seeker_accounting.modules.audit.dto.audit_event_dto import RecordAuditEventCommand
from seeker_accounting.modules.audit.services.audit_service import AuditService
from seeker_accounting.modules.payroll.dto.payroll_export_dto import (
    PayslipBatchExportResultDTO,
    PayslipExportResultDTO,
    SummaryExportResultDTO,
)
from seeker_accounting.modules.payroll.dto.payroll_print_dto import (
    PayrollSummaryPrintDataDTO,
    PayslipPrintDataDTO,
)
from seeker_accounting.modules.payroll.payroll_permissions import PAYROLL_PRINT
from seeker_accounting.modules.payroll.services.payroll_print_service import PayrollPrintService
from seeker_accounting.platform.exceptions import ValidationError


def _safe_filename(raw: str) -> str:
    """Sanitise a string for use in filenames — keep only alphanumeric, dash, underscore."""
    return re.sub(r"[^\w\-]", "_", raw).strip("_")[:80]


# ── Professional HTML/CSS for PDF rendering ──────────────────────────────────
_SUMMARY_CSS = """\
body { font-family: 'Segoe UI', Arial, Helvetica, sans-serif; font-size: 10pt; color: #1F2933; margin: 20px; }
h1 { font-size: 14pt; color: #2F4F6F; margin: 0 0 2px 0; }
.subtitle { font-size: 10pt; color: #6B7280; margin-bottom: 14px; }
table { border-collapse: collapse; width: 100%; font-size: 9.5pt; margin-bottom: 10px; }
th { background-color: #2F4F6F; color: #fff; padding: 6px 10px; font-size: 8.5pt; font-weight: 600; text-align: left; }
th.num { text-align: right; }
td { padding: 4px 10px; border-bottom: 1px solid #EAF1F7; }
td.num { text-align: right; font-variant-numeric: tabular-nums; }
tr.even { background: #ffffff; }
tr.odd { background: #F6F8FB; }
.total-row td { font-weight: 700; border-top: 2px solid #2F4F6F; background: #EAF1F7; color: #2F4F6F; }
.summary-box { margin-top: 14px; }
.summary-box td { padding: 4px 10px; }
.summary-box td.num { font-weight: 600; }
.footer { font-size: 7.5pt; color: #9CA3AF; margin-top: 14px; text-align: right; }
.warning-bar { padding: 6px 10px; background: #fff8e1; border-left: 3px solid #f9a825; font-size: 8.5pt; color: #6d4c00; margin-bottom: 10px; }
"""


def _h(text: str) -> str:
    """HTML-escape a string for safe insertion into HTML output."""
    return html.escape(str(text))


def _fmt(value: Decimal) -> str:
    return f"{value:,.2f}"


class PayrollExportService:
    """Orchestrates payroll PDF and CSV exports from existing print data."""

    def __init__(
        self,
        print_service: PayrollPrintService,
        permission_service: PermissionService,
        audit_service: AuditService,
        logo_resolver: Callable[[str | None], str | None] | None = None,
    ) -> None:
        self._print_service = print_service
        self._permission_service = permission_service
        self._audit_service = audit_service
        self._logo_resolver = logo_resolver

    # ── Individual payslip PDF ────────────────────────────────────────────────

    def export_payslip_pdf(
        self,
        company_id: int,
        run_employee_id: int,
        output_path: str,
        *,
        warning_lines: list[str] | None = None,
    ) -> PayslipExportResultDTO:
        """Export a single payslip to PDF. Returns result DTO."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        data = self._print_service.get_payslip_data(company_id, run_employee_id)
        html_content = self._build_payslip_html_web(data, warning_lines=warning_lines)
        self._render_pdf_via_web(html_content, output_path)

        self._audit_service.record_event(
            company_id,
            RecordAuditEventCommand(
                event_type_code="payslip_exported",
                module_code="payroll",
                entity_type="payroll_run_employee",
                entity_id=run_employee_id,
                description=(
                    f"Payslip exported to PDF for {data.employee_display_name} "
                    f"({data.employee_number}), run {data.run_reference}."
                ),
            ),
        )

        return PayslipExportResultDTO(
            file_path=output_path,
            employee_number=data.employee_number,
            employee_display_name=data.employee_display_name,
            run_reference=data.run_reference,
            period_label=data.period_label,
        )

    # ── Batch payslip PDF ─────────────────────────────────────────────────────

    def export_payslip_batch_pdf(
        self,
        company_id: int,
        run_id: int,
        output_directory: str,
        *,
        run_employee_ids: tuple[int, ...] | None = None,
        warning_lines: list[str] | None = None,
    ) -> PayslipBatchExportResultDTO:
        """Export payslips for all (or selected) employees in a run to individual PDFs."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        payslips = self._print_service.get_payslip_batch_data(
            company_id, run_id, run_employee_ids
        )

        os.makedirs(output_directory, exist_ok=True)

        exported: list[PayslipExportResultDTO] = []
        failed: list[tuple[str, str]] = []

        for data in payslips:
            filename = (
                f"payslip_{_safe_filename(data.run_reference)}"
                f"_{_safe_filename(data.employee_number)}.pdf"
            )
            file_path = os.path.join(output_directory, filename)
            try:
                html_content = self._build_payslip_html_web(data, warning_lines=warning_lines)
                self._render_pdf_via_web(html_content, file_path)
                exported.append(PayslipExportResultDTO(
                    file_path=file_path,
                    employee_number=data.employee_number,
                    employee_display_name=data.employee_display_name,
                    run_reference=data.run_reference,
                    period_label=data.period_label,
                ))
            except Exception as exc:
                failed.append((data.employee_display_name, str(exc)))

        if exported:
            self._audit_service.record_event(
                company_id,
                RecordAuditEventCommand(
                    event_type_code="payslip_batch_exported",
                    module_code="payroll",
                    entity_type="payroll_run",
                    entity_id=run_id,
                    description=(
                        f"Batch payslip PDF export: {len(exported)} exported, "
                        f"{len(failed)} failed."
                    ),
                ),
            )

        return PayslipBatchExportResultDTO(
            output_directory=output_directory,
            exported=tuple(exported),
            failed=tuple(failed),
        )

    # ── Payroll summary CSV ───────────────────────────────────────────────────

    def export_summary_csv(
        self,
        company_id: int,
        run_id: int,
        output_path: str,
    ) -> SummaryExportResultDTO:
        """Export payroll run summary to CSV file."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        data = self._print_service.get_summary_data(company_id, run_id)
        self._write_summary_csv(data, output_path)

        self._audit_service.record_event(
            company_id,
            RecordAuditEventCommand(
                event_type_code="payroll_summary_exported",
                module_code="payroll",
                entity_type="payroll_run",
                entity_id=run_id,
                description=(
                    f"Payroll summary exported to CSV for run {data.run_reference}, "
                    f"period {data.period_label}."
                ),
            ),
        )

        return SummaryExportResultDTO(
            file_path=output_path,
            format="csv",
            run_reference=data.run_reference,
            period_label=data.period_label,
            employee_count=data.employee_count,
        )

    # ── Payroll summary PDF ───────────────────────────────────────────────────

    def export_summary_pdf(
        self,
        company_id: int,
        run_id: int,
        output_path: str,
        *,
        warning_lines: list[str] | None = None,
    ) -> SummaryExportResultDTO:
        """Export payroll run summary to PDF file."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        data = self._print_service.get_summary_data(company_id, run_id)
        html_content = self._build_summary_html(data, warning_lines=warning_lines)
        self._render_pdf_via_web(html_content, output_path)

        self._audit_service.record_event(
            company_id,
            RecordAuditEventCommand(
                event_type_code="payroll_summary_exported",
                module_code="payroll",
                entity_type="payroll_run",
                entity_id=run_id,
                description=(
                    f"Payroll summary exported to PDF for run {data.run_reference}, "
                    f"period {data.period_label}."
                ),
            ),
        )

        return SummaryExportResultDTO(
            file_path=output_path,
            format="pdf",
            run_reference=data.run_reference,
            period_label=data.period_label,
            employee_count=data.employee_count,
        )

    # ══════════════════════════════════════════════════════════════════════════
    # HTML builders
    # ══════════════════════════════════════════════════════════════════════════


    def _build_summary_html(
        self,
        data: PayrollSummaryPrintDataDTO,
        *,
        warning_lines: list[str] | None = None,
    ) -> str:
        """Build professional HTML for payroll summary report."""
        parts: list[str] = [
            "<!DOCTYPE html><html><head><meta charset='utf-8'/>",
            f"<style>{_SUMMARY_CSS}</style></head><body>",
        ]

        # Warning bar
        if warning_lines:
            for w in warning_lines:
                parts.append(f'<div class="warning-bar">{_h(w)}</div>')

        # Header
        parts.append(f"<h1>{_h(data.company_name)}</h1>")
        parts.append(
            f'<div class="subtitle">Payroll Summary &mdash; {_h(data.period_label)}'
            f" &nbsp;&bull;&nbsp; Run: {_h(data.run_reference)} ({_h(data.run_label)})"
            f" &nbsp;&bull;&nbsp; Employees: {data.employee_count}"
            f" &nbsp;&bull;&nbsp; {_h(data.currency_code)}</div>"
        )

        # Employee detail table
        parts.append("<table><tr>")
        parts.append('<th>No.</th><th>Name</th><th class="num">Gross</th>')
        parts.append(f'<th class="num">Deductions &amp; Taxes</th><th class="num">Net Pay ({_h(data.currency_code)})</th>')
        parts.append("</tr>")
        for idx, (emp_no, name, gross, ded_tax, net) in enumerate(data.employee_lines):
            row_cls = "even" if idx % 2 == 0 else "odd"
            parts.append(
                f'<tr class="{row_cls}">'
                f"<td>{_h(emp_no)}</td><td>{_h(name)}</td>"
                f'<td class="num">{_fmt(gross)}</td>'
                f'<td class="num">{_fmt(ded_tax)}</td>'
                f'<td class="num">{_fmt(net)}</td></tr>'
            )
        parts.append(
            '<tr class="total-row"><td colspan="2"><b>TOTALS</b></td>'
            f'<td class="num">{_fmt(data.total_gross_earnings)}</td>'
            f'<td class="num">{_fmt(data.total_deductions + data.total_taxes)}</td>'
            f'<td class="num">{_fmt(data.total_net_payable)}</td></tr>'
        )
        parts.append("</table>")

        # Summary box
        parts.append('<table class="summary-box">')
        parts.append(f'<tr><td>Total Gross Earnings</td><td class="num">{_fmt(data.total_gross_earnings)}</td></tr>')
        parts.append(f'<tr><td>Total Employee Deductions</td><td class="num">{_fmt(data.total_deductions)}</td></tr>')
        parts.append(f'<tr><td>Total Taxes</td><td class="num">{_fmt(data.total_taxes)}</td></tr>')
        parts.append(f'<tr style="font-weight:700"><td>Total Net Payable</td><td class="num">{_fmt(data.total_net_payable)}</td></tr>')
        parts.append(f'<tr><td>Total Employer Contributions</td><td class="num">{_fmt(data.total_employer_contributions)}</td></tr>')
        parts.append(f'<tr style="font-weight:700"><td>Total Employer Cost</td><td class="num">{_fmt(data.total_employer_cost)}</td></tr>')
        parts.append("</table>")

        parts.append(f'<div class="footer">Generated by Seeker Accounting</div>')
        parts.append("</body></html>")
        return "".join(parts)

    # ══════════════════════════════════════════════════════════════════════════
    # CSV writer
    # ══════════════════════════════════════════════════════════════════════════

    @staticmethod
    def _write_summary_csv(data: PayrollSummaryPrintDataDTO, output_path: str) -> None:
        """Write structured CSV for payroll summary."""
        with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.writer(f)

            # Header section
            writer.writerow(["Company", data.company_name])
            writer.writerow(["Run Reference", data.run_reference])
            writer.writerow(["Run Label", data.run_label])
            writer.writerow(["Period", data.period_label])
            writer.writerow(["Currency", data.currency_code])
            writer.writerow(["Employees", data.employee_count])
            writer.writerow([])

            # Employee detail
            writer.writerow(["Employee No.", "Employee Name", "Gross Earnings", "Deductions & Taxes", "Net Pay"])
            for emp_no, name, gross, ded_tax, net in data.employee_lines:
                writer.writerow([emp_no, name, str(gross), str(ded_tax), str(net)])
            writer.writerow([])

            # Totals
            writer.writerow(["TOTALS"])
            writer.writerow(["Total Gross Earnings", str(data.total_gross_earnings)])
            writer.writerow(["Total Employee Deductions", str(data.total_deductions)])
            writer.writerow(["Total Taxes", str(data.total_taxes)])
            writer.writerow(["Total Net Payable", str(data.total_net_payable)])
            writer.writerow(["Total Employer Contributions", str(data.total_employer_contributions)])
            writer.writerow(["Total Employer Cost", str(data.total_employer_cost)])

    # ══════════════════════════════════════════════════════════════════════════
    # PDF renderer
    # ══════════════════════════════════════════════════════════════════════════

    def _build_payslip_html_web(self, data: PayslipPrintDataDTO, *, warning_lines: list[str] | None = None) -> str:
        """Build payslip HTML for Chromium rendering via PayslipHtmlBuilder."""
        from seeker_accounting.modules.payroll.services.payroll_payslip_html_builder import PayslipHtmlBuilder
        builder = PayslipHtmlBuilder(logo_resolver=self._logo_resolver)
        return builder.build(data, warning_lines=warning_lines)

    @staticmethod
    def _render_pdf_via_web(html_content: str, output_path: str) -> None:
        """Render HTML to PDF via Chromium (WebDocumentRenderer)."""
        from seeker_accounting.platform.printing.web_renderer import WebDocumentRenderer
        from seeker_accounting.platform.printing.print_data_protocol import PageOrientation, PageSize
        renderer = WebDocumentRenderer()
        ok = renderer.render_pdf(
            html_content,
            output_path,
            page_size=PageSize.A4,
            orientation=PageOrientation.PORTRAIT,
            margin_mm=0,  # @page CSS rule declares margins
        )
        if not ok:
            raise RuntimeError(f"Chromium PDF rendering failed for: {output_path}")

    # ══════════════════════════════════════════════════════════════════════════
    # Unified format dispatch
    # ══════════════════════════════════════════════════════════════════════════

    def export_payslip_to_format(
        self,
        company_id: int,
        run_employee_id: int,
        output_path: str,
        fmt: str,
        *,
        warning_lines: list[str] | None = None,
    ) -> PayslipExportResultDTO:
        """Export a single payslip in the requested format (pdf / word / excel)."""
        if fmt == "pdf":
            return self.export_payslip_pdf(
                company_id, run_employee_id, output_path, warning_lines=warning_lines
            )
        if fmt == "word":
            return self.export_payslip_word(company_id, run_employee_id, output_path)
        if fmt == "excel":
            return self.export_payslip_excel(company_id, run_employee_id, output_path)
        raise ValueError(f"Unsupported export format: {fmt!r}")

    def export_payslip_batch_to_format(
        self,
        company_id: int,
        run_id: int,
        output_path: str,
        fmt: str,
        *,
        run_employee_ids: tuple[int, ...] | None = None,
        warning_lines: list[str] | None = None,
    ) -> PayslipBatchExportResultDTO:
        """Export payslips for all (or selected) employees in a run.

        - PDF:   individual files written to *output_path* (treated as directory)
        - Word:  individual files written to *output_path* (treated as directory)
        - Excel: single multi-sheet workbook written to *output_path* (treated as file)
        """
        if fmt == "pdf":
            return self.export_payslip_batch_pdf(
                company_id, run_id, output_path,
                run_employee_ids=run_employee_ids, warning_lines=warning_lines,
            )
        if fmt == "word":
            return self.export_payslip_batch_word(
                company_id, run_id, output_path, run_employee_ids=run_employee_ids
            )
        if fmt == "excel":
            return self.export_payslip_batch_excel(
                company_id, run_id, output_path, run_employee_ids=run_employee_ids
            )
        raise ValueError(f"Unsupported export format: {fmt!r}")

    # ══════════════════════════════════════════════════════════════════════════
    # Word export — single and batch
    # ══════════════════════════════════════════════════════════════════════════

    def export_payslip_word(
        self,
        company_id: int,
        run_employee_id: int,
        output_path: str,
    ) -> PayslipExportResultDTO:
        """Export a single payslip to a .docx Word file."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        data = self._print_service.get_payslip_data(company_id, run_employee_id)
        self._write_payslip_word(data, output_path)

        self._audit_service.record_event(
            company_id,
            RecordAuditEventCommand(
                event_type_code="payslip_exported",
                module_code="payroll",
                entity_type="payroll_run_employee",
                entity_id=run_employee_id,
                description=(
                    f"Payslip exported to Word for {data.employee_display_name} "
                    f"({data.employee_number}), run {data.run_reference}."
                ),
            ),
        )
        return PayslipExportResultDTO(
            file_path=output_path,
            employee_number=data.employee_number,
            employee_display_name=data.employee_display_name,
            run_reference=data.run_reference,
            period_label=data.period_label,
        )

    def export_payslip_batch_word(
        self,
        company_id: int,
        run_id: int,
        output_directory: str,
        *,
        run_employee_ids: tuple[int, ...] | None = None,
    ) -> PayslipBatchExportResultDTO:
        """Export individual .docx payslips for all (or selected) employees in a run."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        payslips = self._print_service.get_payslip_batch_data(
            company_id, run_id, run_employee_ids
        )
        os.makedirs(output_directory, exist_ok=True)

        exported: list[PayslipExportResultDTO] = []
        failed: list[tuple[str, str]] = []

        for data in payslips:
            filename = (
                f"payslip_{_safe_filename(data.run_reference)}"
                f"_{_safe_filename(data.employee_number)}.docx"
            )
            file_path = os.path.join(output_directory, filename)
            try:
                self._write_payslip_word(data, file_path)
                exported.append(PayslipExportResultDTO(
                    file_path=file_path,
                    employee_number=data.employee_number,
                    employee_display_name=data.employee_display_name,
                    run_reference=data.run_reference,
                    period_label=data.period_label,
                ))
            except Exception as exc:
                failed.append((data.employee_display_name, str(exc)))

        if exported:
            self._audit_service.record_event(
                company_id,
                RecordAuditEventCommand(
                    event_type_code="payslip_batch_exported",
                    module_code="payroll",
                    entity_type="payroll_run",
                    entity_id=run_id,
                    description=(
                        f"Batch payslip Word export: {len(exported)} exported, "
                        f"{len(failed)} failed."
                    ),
                ),
            )

        return PayslipBatchExportResultDTO(
            output_directory=output_directory,
            exported=tuple(exported),
            failed=tuple(failed),
        )

    # ══════════════════════════════════════════════════════════════════════════
    # Excel export — single and batch (multi-sheet workbook)
    # ══════════════════════════════════════════════════════════════════════════

    def export_payslip_excel(
        self,
        company_id: int,
        run_employee_id: int,
        output_path: str,
    ) -> PayslipExportResultDTO:
        """Export a single payslip to a .xlsx Excel file."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        data = self._print_service.get_payslip_data(company_id, run_employee_id)
        self._write_payslip_excel_single(data, output_path)

        self._audit_service.record_event(
            company_id,
            RecordAuditEventCommand(
                event_type_code="payslip_exported",
                module_code="payroll",
                entity_type="payroll_run_employee",
                entity_id=run_employee_id,
                description=(
                    f"Payslip exported to Excel for {data.employee_display_name} "
                    f"({data.employee_number}), run {data.run_reference}."
                ),
            ),
        )
        return PayslipExportResultDTO(
            file_path=output_path,
            employee_number=data.employee_number,
            employee_display_name=data.employee_display_name,
            run_reference=data.run_reference,
            period_label=data.period_label,
        )

    def export_payslip_batch_excel(
        self,
        company_id: int,
        run_id: int,
        output_path: str,
        *,
        run_employee_ids: tuple[int, ...] | None = None,
    ) -> PayslipBatchExportResultDTO:
        """Export all (or selected) payslips to a single multi-sheet .xlsx workbook."""
        self._permission_service.require_permission(PAYROLL_PRINT)
        payslips = self._print_service.get_payslip_batch_data(
            company_id, run_id, run_employee_ids
        )

        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        exported: list[PayslipExportResultDTO] = []
        failed: list[tuple[str, str]] = []

        try:
            import openpyxl
            wb = openpyxl.Workbook()
            wb.remove(wb.active)  # remove default sheet

            for data in payslips:
                sheet_name = _safe_filename(data.employee_number)[:31]  # Excel sheet name limit
                ws = wb.create_sheet(title=sheet_name)
                self._write_payslip_excel_sheet(data, ws)
                exported.append(PayslipExportResultDTO(
                    file_path=output_path,
                    employee_number=data.employee_number,
                    employee_display_name=data.employee_display_name,
                    run_reference=data.run_reference,
                    period_label=data.period_label,
                ))

            wb.save(output_path)

        except Exception as exc:
            for data in payslips:
                if not any(e.employee_number == data.employee_number for e in exported):
                    failed.append((data.employee_display_name, str(exc)))

        if exported:
            self._audit_service.record_event(
                company_id,
                RecordAuditEventCommand(
                    event_type_code="payslip_batch_exported",
                    module_code="payroll",
                    entity_type="payroll_run",
                    entity_id=run_id,
                    description=(
                        f"Batch payslip Excel export: {len(exported)} employees in workbook, "
                        f"{len(failed)} failed."
                    ),
                ),
            )

        return PayslipBatchExportResultDTO(
            output_directory=os.path.dirname(output_path) or ".",
            exported=tuple(exported),
            failed=tuple(failed),
        )

    # ══════════════════════════════════════════════════════════════════════════
    # Word document renderer
    # ══════════════════════════════════════════════════════════════════════════

    @staticmethod
    def _write_payslip_word(ps: PayslipPrintDataDTO, output_path: str) -> None:
        """Build and save a payslip as a Word .docx document."""
        from docx import Document
        from docx.shared import Mm, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_ORIENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # ── Colours ───────────────────────────────────────────────────────────
        C_PRIMARY = RGBColor(0x2F, 0x4F, 0x6F)
        C_SECONDARY = RGBColor(0x6E, 0x85, 0x9B)
        C_LABEL = RGBColor(0x6B, 0x72, 0x80)
        C_MAIN = RGBColor(0x1F, 0x29, 0x33)
        C_NET = RGBColor(0x2E, 0x7D, 0x4F)

        doc = Document()

        # A4 portrait, 18 mm margins
        sec = doc.sections[0]
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width = Mm(210)
        sec.page_height = Mm(297)
        sec.top_margin = Mm(18)
        sec.bottom_margin = Mm(20)
        sec.left_margin = Mm(18)
        sec.right_margin = Mm(18)

        # Remove the default empty paragraph docx adds
        for p in list(doc.paragraphs):
            p._element.getparent().remove(p._element)

        # ── Inline helpers ────────────────────────────────────────────────────

        def _sp(v: float):
            return Pt(v)

        def _no_borders(table) -> None:
            tbl = table._tbl
            tblPr = tbl.find(qn("w:tblPr"))
            if tblPr is None:
                tblPr = OxmlElement("w:tblPr")
                tbl.insert(0, tblPr)
            tblBrd = OxmlElement("w:tblBorders")
            for bname in ("top", "left", "bottom", "right", "insideH", "insideV"):
                b = OxmlElement(f"w:{bname}")
                b.set(qn("w:val"), "none")
                b.set(qn("w:sz"), "0")
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), "auto")
                tblBrd.append(b)
            tblPr.append(tblBrd)

        def _shade_cell(cell, hex_color: str) -> None:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tcPr.append(shd)

        def _shade_para(para, hex_color: str) -> None:
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            pPr.append(shd)

        def _cell_all_borders(cell, hex_color: str = "D6E0EA", sz: str = "4") -> None:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBrd = OxmlElement("w:tcBorders")
            for bname in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{bname}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), sz)
                b.set(qn("w:space"), "0")
                b.set(qn("w:color"), hex_color)
                tcBrd.append(b)
            tcPr.append(tcBrd)

        def _bottom_border_para(para, hex_color: str = "D6E0EA", sz: str = "4") -> None:
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            b = OxmlElement("w:bottom")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), sz)
            b.set(qn("w:space"), "1")
            b.set(qn("w:color"), hex_color)
            pBdr.append(b)
            pPr.append(pBdr)

        def _add_run(para, text: str, *, size: float = 9.5, bold: bool = False,
                     color: RGBColor | None = None, align=None) -> None:
            if align:
                para.alignment = align
            r = para.add_run(text)
            r.font.size = Pt(size)
            r.bold = bold
            if color:
                r.font.color.rgb = color

        def _fmt(v: Decimal) -> str:
            return f"{v:,.2f}"

        def _cell_pad(cell, top: float = 3, left: float = 6, bottom: float = 3,
                      right: float = 6) -> None:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), str(int(val * 20)))  # twentieths of a point
                m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            tcPr.append(tcMar)

        # ══════════════════════════════════════════════════════════════════════
        # 1. Banner: Company name (left) | BULLETIN DE PAIE (right)
        # ══════════════════════════════════════════════════════════════════════
        banner = doc.add_table(rows=1, cols=2)
        _no_borders(banner)
        banner.rows[0].cells[0].width = Mm(130)
        banner.rows[0].cells[1].width = Mm(50)

        p_name = banner.rows[0].cells[0].paragraphs[0]
        _add_run(p_name, ps.company_name, size=15, bold=True, color=C_PRIMARY)
        p_name.paragraph_format.space_after = _sp(2)

        p_title = banner.rows[0].cells[1].paragraphs[0]
        _add_run(p_title, "BULLETIN DE PAIE\nPAYSLIP", size=10, bold=True,
                 color=C_SECONDARY, align=WD_ALIGN_PARAGRAPH.RIGHT)

        # Separator
        p_sep = doc.add_paragraph()
        _bottom_border_para(p_sep, "D6E0EA", "4")
        p_sep.paragraph_format.space_after = _sp(6)
        p_sep.paragraph_format.space_before = _sp(6)

        # ══════════════════════════════════════════════════════════════════════
        # 2. Identity cards: Employer | Employee
        # ══════════════════════════════════════════════════════════════════════
        hire_str = (
            ps.employee_hire_date.strftime("%d/%m/%Y")
            if ps.employee_hire_date else "—"
        )
        _PA_LABELS = {"bank": "Bank Transfer", "cash": "Cash", "petty_cash": "Petty Cash"}

        EMPLOYER_ROWS: list[tuple[str, str | None]] = [
            ("Company", ps.company_name),
            ("Address / Adresse", ps.company_address),
            ("City / Ville", ps.company_city),
            ("Tax ID / NIU", ps.company_tax_identifier),
            ("CNPS Employer No.", ps.company_cnps_employer_number),
            ("Phone / Tél.", ps.company_phone),
        ]
        EMPLOYEE_ROWS: list[tuple[str, str | None]] = [
            ("Name / Nom", ps.employee_display_name),
            ("Employee No. / Matricule", ps.employee_number),
            ("Job Title / Fonction", ps.employee_position),
            ("Department", ps.employee_department),
            ("Tax ID / NIF", ps.employee_nif),
            ("CNPS No.", ps.employee_cnps_number),
            ("Hire Date / Date d'Embauche", hire_str),
        ]

        id_tbl = doc.add_table(rows=1, cols=2)
        _no_borders(id_tbl)

        def _fill_id_card(cell, card_title: str,
                          rows: list[tuple[str, str | None]]) -> None:
            _shade_cell(cell, "FFFFFF")
            _cell_all_borders(cell, "D6E0EA", "4")
            _cell_pad(cell, top=0, left=8, bottom=6, right=6)

            # Title paragraph (shaded EAF1F7 via paragraph shading)
            p_hdr = cell.paragraphs[0]
            p_hdr.paragraph_format.space_after = _sp(3)
            p_hdr.paragraph_format.space_before = _sp(0)
            _shade_para(p_hdr, "EAF1F7")
            _add_run(p_hdr, f"  {card_title}", size=8.5, bold=True, color=C_PRIMARY)

            for label, value in rows:
                p_row = cell.add_paragraph()
                p_row.paragraph_format.space_after = _sp(1)
                p_row.paragraph_format.space_before = _sp(0)
                r_lbl = p_row.add_run(f"{label}:  ")
                r_lbl.font.size = Pt(7.5)
                r_lbl.font.color.rgb = C_LABEL
                r_val = p_row.add_run(value if value else "—")
                r_val.font.size = Pt(8.5)
                r_val.bold = True
                r_val.font.color.rgb = C_MAIN

        _fill_id_card(id_tbl.rows[0].cells[0], "EMPLOYER / EMPLOYEUR", EMPLOYER_ROWS)
        _fill_id_card(id_tbl.rows[0].cells[1], "EMPLOYEE / EMPLOYÉ(E)", EMPLOYEE_ROWS)

        doc.add_paragraph().paragraph_format.space_after = _sp(4)

        # ══════════════════════════════════════════════════════════════════════
        # 3. Context bar: Pay Period | Payment Date | Run Reference | Mode
        # ══════════════════════════════════════════════════════════════════════
        pay_date_str = ps.payment_date.strftime("%d/%m/%Y") if ps.payment_date else "—"
        if ps.payment_account_name:
            mode_label = _PA_LABELS.get(ps.payment_account_type or "", ps.payment_account_type or "")
            mode_value = ps.payment_account_name
            if ps.payment_account_reference:
                mode_value = f"{mode_value} ({ps.payment_account_reference})"
        else:
            mode_label = "Payment Mode"
            mode_value = "—"

        ctx_data = [
            ("Pay Period", ps.period_label or "—"),
            ("Payment Date", pay_date_str),
            ("Run Reference", ps.run_reference or "—"),
            (mode_label, mode_value),
        ]

        ctx_tbl = doc.add_table(rows=2, cols=4)
        _no_borders(ctx_tbl)
        for col_i, (label, value) in enumerate(ctx_data):
            lbl_cell = ctx_tbl.rows[0].cells[col_i]
            val_cell = ctx_tbl.rows[1].cells[col_i]
            for cell in (lbl_cell, val_cell):
                _shade_cell(cell, "EAF1F7")
                _cell_pad(cell, top=4, left=8, bottom=4, right=8)
            p_lbl = lbl_cell.paragraphs[0]
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_lbl, label, size=7.5, color=C_LABEL)
            p_val = val_cell.paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_val, value, size=9, bold=True, color=C_MAIN)

        doc.add_paragraph().paragraph_format.space_after = _sp(6)

        # ══════════════════════════════════════════════════════════════════════
        # Section builder helper
        # ══════════════════════════════════════════════════════════════════════

        def _add_section(title: str,
                         lines: tuple[tuple[str, Decimal], ...],
                         subtotal: Decimal,
                         subtotal_label: str) -> None:
            if not lines:
                return
            # Section header
            p_hdr = doc.add_paragraph()
            p_hdr.paragraph_format.space_before = _sp(8)
            p_hdr.paragraph_format.space_after = _sp(2)
            _bottom_border_para(p_hdr, "2F4F6F", "6")
            _add_run(p_hdr, title, size=9, bold=True, color=C_PRIMARY)

            # Detail table: 2 columns (description | amount)
            tbl = doc.add_table(rows=0, cols=2)
            _no_borders(tbl)

            for idx, (name, amount) in enumerate(lines):
                row = tbl.add_row()
                bg = "FFFFFF" if idx % 2 == 0 else "F6F8FB"
                for cell in row.cells:
                    _shade_cell(cell, bg)
                    _cell_pad(cell, top=2, left=8, bottom=2, right=8)
                p_name = row.cells[0].paragraphs[0]
                _add_run(p_name, name, size=9, color=C_MAIN)
                # Add question-mark prefix for unknown base components
                p_amt = row.cells[1].paragraphs[0]
                p_amt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _add_run(p_amt, _fmt(amount), size=9, color=C_MAIN)

            # Subtotal row
            sub_row = tbl.add_row()
            for cell in sub_row.cells:
                _shade_cell(cell, "EAF1F7")
                _cell_pad(cell, top=3, left=8, bottom=3, right=8)
            p_sub_lbl = sub_row.cells[0].paragraphs[0]
            _add_run(p_sub_lbl, f"  {subtotal_label}", size=9, bold=True, color=C_PRIMARY)
            p_sub_val = sub_row.cells[1].paragraphs[0]
            p_sub_val.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _add_run(p_sub_val, _fmt(subtotal), size=9, bold=True, color=C_PRIMARY)

        # ══════════════════════════════════════════════════════════════════════
        # 4. Earnings section
        # ══════════════════════════════════════════════════════════════════════
        _add_section(
            "EARNINGS / RÉMUNÉRATIONS",
            ps.earnings, ps.gross_earnings,
            "SUBTOTAL / SOUS-TOTAL",
        )

        # ══════════════════════════════════════════════════════════════════════
        # 5. Statutory bases strip
        # ══════════════════════════════════════════════════════════════════════
        doc.add_paragraph().paragraph_format.space_after = _sp(4)

        bases_tbl = doc.add_table(rows=2, cols=3)
        _no_borders(bases_tbl)
        bases_data = [
            ("CNPS BASE", ps.cnps_contributory_base),
            ("TAXABLE BASE (IRPP)", ps.taxable_salary_base),
            ("TDL BASE", ps.tdl_base),
        ]
        for col_i, (lbl, val) in enumerate(bases_data):
            lbl_cell = bases_tbl.rows[0].cells[col_i]
            val_cell = bases_tbl.rows[1].cells[col_i]
            for cell in (lbl_cell, val_cell):
                _shade_cell(cell, "EAF1F7")
                _cell_all_borders(cell, "D6E0EA", "4")
                _cell_pad(cell, top=6, left=10, bottom=6, right=10)
            p_lbl = lbl_cell.paragraphs[0]
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_lbl, lbl, size=7.5, color=C_LABEL)
            p_val = val_cell.paragraphs[0]
            p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_val, _fmt(val), size=11, bold=True, color=C_PRIMARY)

        doc.add_paragraph().paragraph_format.space_after = _sp(4)

        # ══════════════════════════════════════════════════════════════════════
        # 6. Deductions section
        # ══════════════════════════════════════════════════════════════════════
        _add_section(
            "EMPLOYEE DEDUCTIONS / RETENUES SALARIALES",
            ps.deductions, ps.total_deductions,
            "SUBTOTAL / SOUS-TOTAL",
        )

        # ══════════════════════════════════════════════════════════════════════
        # 7. Taxes section
        # ══════════════════════════════════════════════════════════════════════
        _add_section(
            "TAXES / IMPÔTS",
            ps.taxes, ps.total_taxes,
            "TOTAL TAXES",
        )

        # ══════════════════════════════════════════════════════════════════════
        # 8. Net pay box
        # ══════════════════════════════════════════════════════════════════════
        doc.add_paragraph().paragraph_format.space_after = _sp(6)

        net_taxable = ps.taxable_salary_base - ps.total_deductions
        net_tbl = doc.add_table(rows=2, cols=2)
        _no_borders(net_tbl)

        # Row 1: Net Taxable Pay
        for cell in net_tbl.rows[0].cells:
            _shade_cell(cell, "EDF7F1")
            _cell_all_borders(cell, "C3DFD0", "4")
            _cell_pad(cell, top=6, left=18, bottom=4, right=18)
        p_ntl = net_tbl.rows[0].cells[0].paragraphs[0]
        _add_run(p_ntl, "Net Taxable Pay / Salaire Net Imposable", size=9, color=C_NET)
        p_ntv = net_tbl.rows[0].cells[1].paragraphs[0]
        p_ntv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(p_ntv, f"{_fmt(net_taxable)} {ps.currency_code}", size=9.5,
                 bold=True, color=C_NET)

        # Row 2: NET PAYABLE
        for cell in net_tbl.rows[1].cells:
            _shade_cell(cell, "EDF7F1")
            _cell_all_borders(cell, "C3DFD0", "6")
            _cell_pad(cell, top=8, left=18, bottom=8, right=18)
        p_npl = net_tbl.rows[1].cells[0].paragraphs[0]
        _add_run(p_npl, "NET PAYABLE / SALAIRE NET À PAYER", size=12, bold=True,
                 color=C_NET)
        p_npv = net_tbl.rows[1].cells[1].paragraphs[0]
        p_npv.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_run(p_npv, f"{_fmt(ps.net_payable)} {ps.currency_code}", size=17,
                 bold=True, color=C_NET)

        doc.add_paragraph().paragraph_format.space_after = _sp(6)

        # ══════════════════════════════════════════════════════════════════════
        # 9. Employer contributions
        # ══════════════════════════════════════════════════════════════════════
        _add_section(
            "EMPLOYER CHARGES / CHARGES PATRONALES",
            ps.employer_contributions, ps.total_employer_contributions,
            "TOTAL EMPLOYER CHARGES",
        )

        # ══════════════════════════════════════════════════════════════════════
        # 10. Signature block
        # ══════════════════════════════════════════════════════════════════════
        doc.add_paragraph().paragraph_format.space_after = _sp(14)
        p_sig_sep = doc.add_paragraph()
        _bottom_border_para(p_sig_sep, "D6E0EA", "4")
        p_sig_sep.paragraph_format.space_after = _sp(4)

        sig_tbl = doc.add_table(rows=1, cols=3)
        _no_borders(sig_tbl)
        for col_i, title in enumerate((
            "Prepared by / Établi par",
            "Approved by / Approuvé par",
            "Employee / Employé(e)",
        )):
            cell = sig_tbl.rows[0].cells[col_i]
            _cell_pad(cell, top=0, left=8, bottom=0, right=8)
            # Sig space
            p_sp = cell.paragraphs[0]
            p_sp.paragraph_format.space_before = _sp(36)
            p_sp.paragraph_format.space_after = _sp(0)
            # Sig line
            p_line = cell.add_paragraph()
            _bottom_border_para(p_line, "6E859B", "4")
            p_line.paragraph_format.space_after = _sp(0)
            # Sig label
            p_lbl = cell.add_paragraph()
            p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_run(p_lbl, title, size=8, color=C_LABEL)

        # Footer note
        doc.add_paragraph().paragraph_format.space_after = _sp(4)
        p_footer = doc.add_paragraph()
        _add_run(p_footer, "Generated by Seeker Accounting", size=7.5, color=C_LABEL)
        p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Save
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        doc.save(output_path)

    # ══════════════════════════════════════════════════════════════════════════
    # Excel renderer
    # ══════════════════════════════════════════════════════════════════════════

    @staticmethod
    def _write_payslip_excel_sheet(ps: PayslipPrintDataDTO, ws) -> None:
        """Write a complete payslip onto an openpyxl worksheet (ws)."""
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
        from openpyxl.utils import get_column_letter

        # ── Colours ───────────────────────────────────────────────────────────
        C_PRIMARY  = "2F4F6F"
        C_SECTION  = "EAF1F7"
        C_CARD_BDR = "D6E0EA"
        C_LABEL    = "6B7280"
        C_MAIN     = "1F2933"
        C_NET      = "2E7D4F"
        C_NET_BG   = "EDF7F1"
        C_WHITE    = "FFFFFF"
        C_STRIPE   = "F6F8FB"

        def _font(size=9, bold=False, color=C_MAIN, italic=False):
            return Font(name="Calibri", size=size, bold=bold, color=color, italic=italic)

        def _fill(color):
            return PatternFill(fill_type="solid", fgColor=color)

        def _align(h="left", v="center", wrap=False):
            return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

        def _border_thin(color=C_CARD_BDR):
            s = Side(border_style="thin", color=color)
            return Border(left=s, right=s, top=s, bottom=s)

        def _border_bottom(color=C_CARD_BDR):
            s = Side(border_style="thin", color=color)
            return Border(bottom=s)

        def _fmt(v: Decimal) -> float:
            return float(v)

        # Column widths: A=wide description, B=amounts
        ws.column_dimensions["A"].width = 48
        ws.column_dimensions["B"].width = 18
        ws.column_dimensions["C"].width = 28
        ws.column_dimensions["D"].width = 18

        # Page setup: A4 portrait
        ws.page_setup.paperSize = 9
        ws.page_setup.orientation = "portrait"
        ws.page_setup.fitToPage = True
        ws.page_setup.fitToWidth = 1

        row = 1

        def _w(col, value, *, font=None, fill=None, alignment=None, border=None,
               number_format=None):
            nonlocal row
            cell = ws.cell(row=row, column=col, value=value)
            if font:
                cell.font = font
            if fill:
                cell.fill = fill
            if alignment:
                cell.alignment = alignment
            if border:
                cell.border = border
            if number_format:
                cell.number_format = number_format
            return cell

        # ══════════════════════════════════════════════════════════════════════
        # Banner row
        # ══════════════════════════════════════════════════════════════════════
        _w(1, ps.company_name, font=_font(14, bold=True, color=C_PRIMARY),
           alignment=_align("left", "center"))
        _w(3, "BULLETIN DE PAIE / PAYSLIP",
           font=_font(10, bold=True, color="6E859B"),
           alignment=_align("right", "center"))
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
        ws.row_dimensions[row].height = 22
        row += 1

        # ══════════════════════════════════════════════════════════════════════
        # Identity info block: Employer (A-B) | Employee (C-D)
        # ══════════════════════════════════════════════════════════════════════
        row += 1  # blank

        hire_str = (
            ps.employee_hire_date.strftime("%d/%m/%Y")
            if ps.employee_hire_date else "—"
        )
        _PA_LABELS = {"bank": "Bank Transfer", "cash": "Cash", "petty_cash": "Petty Cash"}

        EMPLOYER_ROWS: list[tuple[str, str | None]] = [
            ("EMPLOYER / EMPLOYEUR", None),
            ("Company", ps.company_name),
            ("Address / Adresse", ps.company_address),
            ("City / Ville", ps.company_city),
            ("Tax ID / NIU", ps.company_tax_identifier),
            ("CNPS Employer No.", ps.company_cnps_employer_number),
            ("Phone / Tél.", ps.company_phone),
        ]
        EMPLOYEE_ROWS: list[tuple[str, str | None]] = [
            ("EMPLOYEE / EMPLOYÉ(E)", None),
            ("Name / Nom", ps.employee_display_name),
            ("Employee No. / Matricule", ps.employee_number),
            ("Job Title / Fonction", ps.employee_position),
            ("Department", ps.employee_department),
            ("Tax ID / NIF", ps.employee_nif),
            ("CNPS No.", ps.employee_cnps_number),
            ("Hire Date / Date d'Embauche", hire_str),
        ]

        id_start_row = row
        for i, ((el, ev), (rl, rv)) in enumerate(zip(EMPLOYER_ROWS, EMPLOYEE_ROWS)):
            is_hdr = i == 0
            bg = C_SECTION if is_hdr else C_WHITE
            bdr = _border_thin()

            c1 = ws.cell(row=row, column=1, value=el)
            c2 = ws.cell(row=row, column=2, value=ev or ("—" if not is_hdr else None))
            c3 = ws.cell(row=row, column=3, value=rl)
            c4 = ws.cell(row=row, column=4, value=rv or ("—" if not is_hdr else None))

            for cell in (c1, c2, c3, c4):
                cell.fill = _fill(bg)
                cell.border = bdr

            if is_hdr:
                for cell in (c1, c3):
                    cell.font = _font(8.5, bold=True, color=C_PRIMARY)
                    cell.alignment = _align("left", "center")
                ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
                ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
            else:
                c1.font = _font(7.5, color=C_LABEL)
                c1.alignment = _align("right", "center")
                c2.font = _font(8.5, bold=True, color=C_MAIN)
                c2.alignment = _align("left", "center")
                c3.font = _font(7.5, color=C_LABEL)
                c3.alignment = _align("right", "center")
                c4.font = _font(8.5, bold=True, color=C_MAIN)
                c4.alignment = _align("left", "center")

            ws.row_dimensions[row].height = 16
            row += 1

        row += 1  # blank

        # ══════════════════════════════════════════════════════════════════════
        # Context bar: Pay Period | Payment Date | Run Reference | Mode
        # ══════════════════════════════════════════════════════════════════════
        pay_date_str = ps.payment_date.strftime("%d/%m/%Y") if ps.payment_date else "—"
        if ps.payment_account_name:
            mode_label = _PA_LABELS.get(ps.payment_account_type or "", ps.payment_account_type or "")
            mode_value = ps.payment_account_name
            if ps.payment_account_reference:
                mode_value = f"{mode_value} ({ps.payment_account_reference})"
        else:
            mode_label = "Payment Mode"
            mode_value = "—"

        ctx_labels = ["Pay Period", "Payment Date", "Run Reference", mode_label]
        ctx_values = [ps.period_label, pay_date_str, ps.run_reference or "—", mode_value]

        for li, vi in zip(ctx_labels, ctx_values):
            c = ws.cell(row=row, column=ctx_labels.index(li) + 1, value=li)
            c.font = _font(7.5, color=C_LABEL)
            c.fill = _fill(C_SECTION)
            c.alignment = _align("center", "center")
            c.border = _border_thin()
        ws.row_dimensions[row].height = 14
        row += 1

        for li, vi in zip(ctx_labels, ctx_values):
            c = ws.cell(row=row, column=ctx_labels.index(li) + 1, value=vi)
            c.font = _font(9, bold=True, color=C_MAIN)
            c.fill = _fill(C_SECTION)
            c.alignment = _align("center", "center")
            c.border = _border_thin()
        ws.row_dimensions[row].height = 16
        row += 2  # blank

        # ══════════════════════════════════════════════════════════════════════
        # Section builder
        # ══════════════════════════════════════════════════════════════════════

        def _write_section(title: str,
                           lines: tuple[tuple[str, Decimal], ...],
                           subtotal: Decimal,
                           subtotal_label: str) -> None:
            nonlocal row
            if not lines:
                return

            # Section header (spans A-B)
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
            c = ws.cell(row=row, column=1, value=title)
            c.font = _font(9, bold=True, color=C_PRIMARY)
            c.fill = _fill(C_SECTION)
            c.alignment = _align("left", "center")
            c.border = _border_bottom("2F4F6F")
            ws.row_dimensions[row].height = 16
            row += 1

            for idx, (name, amount) in enumerate(lines):
                bg = C_WHITE if idx % 2 == 0 else C_STRIPE
                cn = ws.cell(row=row, column=1, value=name)
                ca = ws.cell(row=row, column=2, value=_fmt(amount))
                cn.font = _font(9, color=C_MAIN)
                cn.fill = _fill(bg)
                cn.alignment = _align("left", "center")
                ca.font = _font(9, color=C_MAIN)
                ca.fill = _fill(bg)
                ca.alignment = _align("right", "center")
                ca.number_format = '#,##0.00'
                ws.row_dimensions[row].height = 15
                row += 1

            # Subtotal
            ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=1)
            cs = ws.cell(row=row, column=1, value=f"  {subtotal_label}")
            csv_ = ws.cell(row=row, column=2, value=_fmt(subtotal))
            cs.font = _font(9, bold=True, color=C_PRIMARY)
            cs.fill = _fill(C_SECTION)
            cs.border = _border_thin(C_CARD_BDR)
            cs.alignment = _align("left", "center")
            csv_.font = _font(9, bold=True, color=C_PRIMARY)
            csv_.fill = _fill(C_SECTION)
            csv_.border = _border_thin(C_CARD_BDR)
            csv_.alignment = _align("right", "center")
            csv_.number_format = '#,##0.00'
            ws.row_dimensions[row].height = 16
            row += 2  # blank after section

        _write_section(
            "EARNINGS / RÉMUNÉRATIONS",
            ps.earnings, ps.gross_earnings, "SUBTOTAL / SOUS-TOTAL",
        )

        # ══════════════════════════════════════════════════════════════════════
        # Statutory bases
        # ══════════════════════════════════════════════════════════════════════
        bases = [
            ("CNPS BASE", ps.cnps_contributory_base),
            ("TAXABLE BASE (IRPP)", ps.taxable_salary_base),
            ("TDL BASE", ps.tdl_base),
        ]
        for col_i, (lbl, val) in enumerate(bases):
            cl = ws.cell(row=row, column=col_i * 1 + 1 if col_i < 2 else 4 - 1,
                         value=lbl)
            # Simple: write across A, B, C as label rows for bases
        # Simpler flat approach: two rows (labels | values) across columns A, B, C
        for col_i, (lbl, val) in enumerate(bases):
            c_lbl = ws.cell(row=row, column=col_i + 1, value=lbl)
            c_lbl.font = _font(7.5, color=C_LABEL)
            c_lbl.fill = _fill(C_SECTION)
            c_lbl.alignment = _align("center", "center")
            c_lbl.border = _border_thin()
        ws.row_dimensions[row].height = 14
        row += 1

        for col_i, (lbl, val) in enumerate(bases):
            c_val = ws.cell(row=row, column=col_i + 1, value=_fmt(val))
            c_val.font = _font(11, bold=True, color=C_PRIMARY)
            c_val.fill = _fill(C_SECTION)
            c_val.alignment = _align("center", "center")
            c_val.border = _border_thin()
            c_val.number_format = '#,##0.00'
        ws.row_dimensions[row].height = 18
        row += 2  # blank

        _write_section(
            "EMPLOYEE DEDUCTIONS / RETENUES SALARIALES",
            ps.deductions, ps.total_deductions, "SUBTOTAL / SOUS-TOTAL",
        )

        _write_section(
            "TAXES / IMPÔTS",
            ps.taxes, ps.total_taxes, "TOTAL TAXES",
        )

        # ══════════════════════════════════════════════════════════════════════
        # Net pay rows
        # ══════════════════════════════════════════════════════════════════════
        net_taxable = ps.taxable_salary_base - ps.total_deductions

        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=1)
        c1 = ws.cell(row=row, column=1, value="Net Taxable Pay / Salaire Net Imposable")
        c2 = ws.cell(row=row, column=2, value=_fmt(net_taxable))
        for c in (c1, c2):
            c.fill = _fill(C_NET_BG)
            c.border = _border_thin("C3DFD0")
            c.font = _font(9, color=C_NET)
        c2.alignment = _align("right", "center")
        c2.number_format = '#,##0.00'
        ws.row_dimensions[row].height = 16
        row += 1

        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=1)
        c1 = ws.cell(row=row, column=1, value="NET PAYABLE / SALAIRE NET À PAYER")
        c2 = ws.cell(row=row, column=2, value=_fmt(ps.net_payable))
        for c in (c1, c2):
            c.fill = _fill(C_NET_BG)
            c.border = _border_thin("C3DFD0")
        c1.font = _font(12, bold=True, color=C_NET)
        c1.alignment = _align("left", "center")
        c2.font = _font(16, bold=True, color=C_NET)
        c2.alignment = _align("right", "center")
        c2.number_format = '#,##0.00'
        ws.row_dimensions[row].height = 22
        row += 2  # blank

        _write_section(
            "EMPLOYER CHARGES / CHARGES PATRONALES",
            ps.employer_contributions, ps.total_employer_contributions,
            "TOTAL EMPLOYER CHARGES",
        )

        # ══════════════════════════════════════════════════════════════════════
        # Signature block (text only, no actual lines in Excel)
        # ══════════════════════════════════════════════════════════════════════
        for col_i, title in enumerate((
            "Prepared by / Établi par",
            "Approved by / Approuvé par",
            "Employee / Employé(e)",
        )):
            c = ws.cell(row=row, column=col_i + 1, value=title)
            c.font = _font(8, color=C_LABEL)
            c.fill = _fill(C_WHITE)
            c.alignment = _align("center", "center")
            c.border = _border_bottom("6E859B")
        ws.row_dimensions[row].height = 30
        row += 1

        # Footer
        c_ft = ws.cell(row=row + 1, column=1, value="Generated by Seeker Accounting")
        c_ft.font = _font(7.5, italic=True, color=C_LABEL)

    @staticmethod
    def _write_payslip_excel_single(ps: PayslipPrintDataDTO, output_path: str) -> None:
        """Write a payslip to a single-sheet .xlsx workbook."""
        import openpyxl
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = _safe_filename(ps.employee_number)[:31]
        PayrollExportService._write_payslip_excel_sheet(ps, ws)
        wb.save(output_path)
