"""python-docx based Word document builder for Seeker Accounting exports.

Produces .docx files with:
  - Proper A4 or A5 page setup with margins
  - Company identity header
  - Seeker Accounting branded footer
  - Professional table styling using python-docx's native API
  - Section headings and paragraph formatting

Usage:
    builder = WordDocumentBuilder(page_size=PageSize.A4)
    builder.add_company_header(company_data)
    builder.add_document_title("Sales Invoice", subtitle="INV-0042")
    builder.add_key_value_grid([("Date", "2026-03-31"), ("Customer", "Acme Corp")])
    builder.add_section_title("Line Items")
    builder.add_data_table(columns, rows, numeric_columns={2, 3, 4})
    builder.add_totals_pairs([("Total", "1,500,000")])
    builder.save("/path/to/output.docx")
"""
from __future__ import annotations

import os
from typing import TYPE_CHECKING

from seeker_accounting.platform.printing.print_data_protocol import (
    CompanyHeaderData,
    PageOrientation,
    PageSize,
)

if TYPE_CHECKING:
    pass

# ── Brand constants ─────────────────────────────────────────────────────────────

_BRAND_NAME = "Seeker Accounting"
_BRAND_TAGLINE = "Built for Business Clarity. Designed For Success."

# RGB tuples for python-docx RGBColor
_RGB_BRAND_PRIMARY = (30, 58, 95)     # #1e3a5f
_RGB_BRAND_LIGHT = (240, 243, 247)    # #f0f3f7
_RGB_MUTED = (107, 114, 128)          # #6b7280
_RGB_WHITE = (255, 255, 255)
_RGB_TOTAL_BG = (238, 242, 247)       # #eef2f7
_RGB_STRIPE = (249, 250, 251)         # #f9fafb


def _mm(value: float):  # type: ignore[return]
    """Convert mm to python-docx Mm units (lazy import)."""
    from docx.shared import Mm
    return Mm(value)


def _pt(value: float):  # type: ignore[return]
    """Convert pt to python-docx Pt units (lazy import)."""
    from docx.shared import Pt
    return Pt(value)


def _rgb(r: int, g: int, b: int):  # type: ignore[return]
    """Create an RGBColor (lazy import)."""
    from docx.shared import RGBColor
    return RGBColor(r, g, b)


class WordDocumentBuilder:
    """Builds a Word (.docx) document with Seeker Accounting's visual language."""

    def __init__(
        self,
        page_size: PageSize = PageSize.A4,
        orientation: PageOrientation = PageOrientation.PORTRAIT,
    ) -> None:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        self._doc = Document()
        self._page_size = page_size
        self._orientation = orientation
        self._WD_ALIGN = WD_ALIGN_PARAGRAPH

        # Page setup
        section = self._doc.sections[0]
        w_mm, h_mm = page_size.dimensions_mm
        if orientation == PageOrientation.LANDSCAPE:
            w_mm, h_mm = h_mm, w_mm

        section.page_width = _mm(w_mm)
        section.page_height = _mm(h_mm)
        section.orientation = (
            WD_ORIENT.LANDSCAPE if orientation == PageOrientation.LANDSCAPE
            else WD_ORIENT.PORTRAIT
        )

        margin = page_size.margin_mm
        section.top_margin = _mm(margin)
        section.bottom_margin = _mm(margin + 12)  # extra space for footer
        section.left_margin = _mm(margin)
        section.right_margin = _mm(margin)

        # Remove default empty paragraph
        for elem in list(self._doc.paragraphs):
            elem._element.getparent().remove(elem._element)  # type: ignore[attr-defined]

        self._add_word_footer(section)

    # ── Public API ──────────────────────────────────────────────────────────────

    def add_company_header(self, company: CompanyHeaderData) -> "WordDocumentBuilder":
        """Add company identity block at the top of the document."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # Company name (large, brand colour)
        p = self._doc.add_paragraph()
        run = p.add_run(company.name)
        run.bold = True
        run.font.size = _pt(self._page_size.body_font_size_pt + 4)
        run.font.color.rgb = _rgb(*_RGB_BRAND_PRIMARY)
        p.paragraph_format.space_after = _pt(0)

        if company.legal_name and company.legal_name != company.name:
            p2 = self._doc.add_paragraph()
            run2 = p2.add_run(company.legal_name)
            run2.font.size = _pt(self._page_size.body_font_size_pt - 1)
            run2.font.color.rgb = _rgb(*_RGB_MUTED)
            p2.paragraph_format.space_after = _pt(0)

        # Address lines
        for line in company.address_block:
            ad = self._doc.add_paragraph()
            adr = ad.add_run(line)
            adr.font.size = _pt(self._page_size.body_font_size_pt - 1)
            ad.paragraph_format.space_after = _pt(0)

        # Contact / tax meta
        meta: list[str] = []
        if company.tax_identifier:
            meta.append(f"Tax ID: {company.tax_identifier}")
        if company.registration_number:
            meta.append(f"Reg: {company.registration_number}")
        if company.phone:
            meta.append(company.phone)
        if company.email:
            meta.append(company.email)
        if meta:
            pm = self._doc.add_paragraph()
            mr = pm.add_run("  ·  ".join(meta))
            mr.font.size = _pt(self._page_size.body_font_size_pt - 1)
            mr.font.color.rgb = _rgb(*_RGB_MUTED)
            pm.paragraph_format.space_after = _pt(0)

        self._add_ruled_separator()
        return self

    def add_document_title(
        self,
        title: str,
        *,
        subtitle: str | None = None,
    ) -> "WordDocumentBuilder":
        """Add the main document title (e.g. 'Sales Invoice — INV-0042')."""
        p = self._doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = _pt(self._page_size.body_font_size_pt + 2)
        r.font.color.rgb = _rgb(*_RGB_BRAND_PRIMARY)
        p.paragraph_format.space_after = _pt(2)
        p.paragraph_format.space_before = _pt(4)

        if subtitle:
            ps = self._doc.add_paragraph()
            rs = ps.add_run(subtitle)
            rs.font.size = _pt(self._page_size.body_font_size_pt - 1)
            rs.font.color.rgb = _rgb(*_RGB_MUTED)
            ps.paragraph_format.space_after = _pt(4)

        return self

    def add_key_value_grid(
        self,
        pairs: list[tuple[str, str]],
        columns: int = 3,
    ) -> "WordDocumentBuilder":
        """Add a key-value metadata grid (inline label: value pairs)."""
        per_row = columns
        col_widths = [self._usable_width_mm() / (per_row * 2)] * (per_row * 2)

        for i in range(0, len(pairs), per_row):
            chunk = pairs[i : i + per_row]
            # Pad to fill the row
            while len(chunk) < per_row:
                chunk.append(("", ""))

            table = self._doc.add_table(rows=1, cols=per_row * 2)
            table.style = "Table Grid"
            for border_el in table._tbl.iter():  # remove borders
                break

            row = table.rows[0]
            for col_idx, (label, value) in enumerate(chunk):
                lc = row.cells[col_idx * 2]
                vc = row.cells[col_idx * 2 + 1]
                lc.width = _mm(self._page_size.margin_mm * 2.5)
                vc.width = _mm(self._usable_width_mm() / per_row - self._page_size.margin_mm * 2.5)

                lp = lc.paragraphs[0]
                lr = lp.add_run(label.upper())
                lr.font.size = _pt(self._page_size.body_font_size_pt - 2)
                lr.font.color.rgb = _rgb(*_RGB_MUTED)

                vp = vc.paragraphs[0]
                vr = vp.add_run(value)
                vr.bold = True
                vr.font.size = _pt(self._page_size.body_font_size_pt - 1)

            self._remove_table_borders(table)
            self._add_space_after_paragraph(_pt(2))

        return self

    def add_section_title(self, title: str) -> "WordDocumentBuilder":
        """Add a section heading (e.g. 'Line Items')."""
        p = self._doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = _pt(self._page_size.body_font_size_pt)
        r.font.color.rgb = _rgb(*_RGB_BRAND_PRIMARY)
        p.paragraph_format.space_before = _pt(8)
        p.paragraph_format.space_after = _pt(3)
        self._add_bottom_border_to_paragraph(p)
        return self

    def add_data_table(
        self,
        columns: list[str],
        rows: list[list[str]],
        *,
        numeric_columns: set[int] | None = None,
        total_row: list[str] | None = None,
    ) -> "WordDocumentBuilder":
        """Add a styled data table."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        num_cols = numeric_columns or set()
        all_rows = rows + ([total_row] if total_row else [])
        table = self._doc.add_table(rows=1 + len(all_rows), cols=len(columns))

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(columns):
            cell = hdr_cells[i]
            self._shade_cell(cell, _RGB_BRAND_PRIMARY)
            p = cell.paragraphs[0]
            run = p.add_run(col_name)
            run.bold = True
            run.font.size = _pt(self._page_size.body_font_size_pt - 1)
            run.font.color.rgb = _rgb(*_RGB_WHITE)
            if i in num_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Data rows
        for row_idx, row_data in enumerate(all_rows):
            is_total = total_row is not None and row_idx == len(rows)
            cells = table.rows[row_idx + 1].cells
            bg = _RGB_TOTAL_BG if is_total else (_RGB_WHITE if row_idx % 2 == 0 else _RGB_STRIPE)
            for col_idx, cell_val in enumerate(row_data):
                cell = cells[col_idx]
                self._shade_cell(cell, bg)
                p = cell.paragraphs[0]
                run = p.add_run(cell_val)
                run.font.size = _pt(self._page_size.body_font_size_pt - 1)
                if is_total:
                    run.bold = True
                    run.font.color.rgb = _rgb(*_RGB_BRAND_PRIMARY)
                if col_idx in num_cols:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        self._doc.add_paragraph()
        return self

    def add_summary_pairs(
        self,
        pairs: list[tuple[str, str]],
        *,
        highlight_last: bool = True,
    ) -> "WordDocumentBuilder":
        """Add a right-aligned totals/summary table (label + value pairs)."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for i, (label, value) in enumerate(pairs):
            is_last = highlight_last and i == len(pairs) - 1
            table = self._doc.add_table(rows=1, cols=2)
            row = table.rows[0]
            bg = _RGB_TOTAL_BG if is_last else _RGB_WHITE

            lc, vc = row.cells[0], row.cells[1]
            self._shade_cell(lc, bg)
            self._shade_cell(vc, bg)

            lp, vp = lc.paragraphs[0], vc.paragraphs[0]
            lr = lp.add_run(label)
            vr = vp.add_run(value)
            fs = self._page_size.body_font_size_pt - 1
            lr.font.size = _pt(fs)
            vr.font.size = _pt(fs)
            if is_last:
                lr.bold = True
                vr.bold = True
                lr.font.color.rgb = _rgb(*_RGB_BRAND_PRIMARY)
                vr.font.color.rgb = _rgb(*_RGB_BRAND_PRIMARY)
            vp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            self._remove_table_borders(table)

        return self

    def add_paragraph(self, text: str, *, italic: bool = False) -> "WordDocumentBuilder":
        """Add a simple body paragraph."""
        p = self._doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = _pt(self._page_size.body_font_size_pt - 1)
        if italic:
            r.italic = True
            r.font.color.rgb = _rgb(*_RGB_MUTED)
        return self

    def add_page_break(self) -> "WordDocumentBuilder":
        """Insert a manual page break."""
        self._doc.add_page_break()
        return self

    def save(self, output_path: str) -> None:
        """Save the document to the given path."""
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
        self._doc.save(output_path)

    # ── Internal helpers ────────────────────────────────────────────────────────

    def _usable_width_mm(self) -> float:
        w_mm, _ = self._page_size.dimensions_mm
        if self._orientation == PageOrientation.LANDSCAPE:
            _, w_mm = self._page_size.dimensions_mm
        return w_mm - (self._page_size.margin_mm * 2)

    def _add_ruled_separator(self) -> None:
        """Add a bottom-border paragraph as a visual separator."""
        p = self._doc.add_paragraph()
        p.paragraph_format.space_after = _pt(6)
        self._add_bottom_border_to_paragraph(p, width="6", color="1E3A5F")

    def _add_space_after_paragraph(self, space) -> None:  # type: ignore[type-arg]
        if self._doc.paragraphs:
            self._doc.paragraphs[-1].paragraph_format.space_after = space

    def _add_word_footer(self, section) -> None:  # type: ignore[type-arg]
        """Add the Seeker Accounting branded footer to this section."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        footer = section.footer
        # Remove existing paragraphs
        for p in footer.paragraphs:
            p.clear()
        if not footer.paragraphs:
            footer._element.append(OxmlElement("w:p"))

        p = footer.paragraphs[0]
        p.paragraph_format.space_before = _pt(0)
        p.paragraph_format.space_after = _pt(0)

        # Left: brand name
        run_brand = p.add_run(_BRAND_NAME)
        run_brand.bold = True
        run_brand.font.size = _pt(7.5)
        run_brand.font.color.rgb = _rgb(*_RGB_BRAND_PRIMARY)

        p.add_run("  ·  ").font.size = _pt(7)

        # Tagline
        run_tagline = p.add_run(_BRAND_TAGLINE)
        run_tagline.italic = True
        run_tagline.font.size = _pt(7)
        run_tagline.font.color.rgb = _rgb(*_RGB_MUTED)

        # Right: page number (tab + field)
        tab_stop_xml = OxmlElement("w:tab")
        p._p.append(tab_stop_xml)  # type: ignore[attr-defined]
        run_page = p.add_run()
        run_page.font.size = _pt(7)
        run_page.font.color.rgb = _rgb(*_RGB_MUTED)
        # Add page number field
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.text = " PAGE "
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        for el in [fld_char1, instr, fld_char2]:
            run_page._r.append(el)  # type: ignore[attr-defined]

    @staticmethod
    def _shade_cell(cell, rgb: tuple[int, int, int]) -> None:  # type: ignore[type-arg]
        """Apply background shading to a table cell."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
        shd = OxmlElement("w:shd")
        hex_color = "{:02X}{:02X}{:02X}".format(*rgb)
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tc_pr.append(shd)

    @staticmethod
    def _remove_table_borders(table) -> None:  # type: ignore[type-arg]
        """Remove all borders from a table."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        tbl_pr = table._tbl.tblPr  # type: ignore[attr-defined]
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)
        tbl_borders = OxmlElement("w:tblBorders")
        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tbl_borders.append(border)
        tbl_pr.append(tbl_borders)

    @staticmethod
    def _add_bottom_border_to_paragraph(p, *, width: str = "4", color: str = "1E3A5F") -> None:
        """Add a bottom border line to a paragraph."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        p_pr = p._p.get_or_add_pPr()  # type: ignore[attr-defined]
        p_borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), width)
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), color)
        p_borders.append(bottom)
        p_pr.append(p_borders)
