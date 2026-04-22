"""Research-only extraction of Cameroon taxation documentation.

Writes text dumps under artifacts/tax_research/ for analysis.
"""
from __future__ import annotations

import os
from pathlib import Path

SRC_DIR = Path(r"C:\Users\User\Downloads\Taxation documentation")
OUT_DIR = Path(r"C:\Users\User\Desktop\Seeker Accounting\artifacts\tax_research")
OUT_DIR.mkdir(parents=True, exist_ok=True)


def extract_docx(path: Path) -> str:
    import docx  # python-docx
    doc = docx.Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for i, table in enumerate(doc.tables):
        parts.append(f"\n[TABLE {i}]")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_pdf(path: Path) -> str:
    import pypdf
    reader = pypdf.PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001
            text = f"[extract error page {i}: {exc}]"
        parts.append(f"\n=== PAGE {i + 1} ===\n{text}")
    return "\n".join(parts)


def safe_name(name: str) -> str:
    return "".join(ch if ch.isalnum() or ch in "._-" else "_" for ch in name)


def main() -> None:
    for path in sorted(SRC_DIR.iterdir()):
        if not path.is_file():
            continue
        suffix = path.suffix.lower()
        out_name = safe_name(path.stem) + ".txt"
        out_path = OUT_DIR / out_name
        if out_path.exists() and out_path.stat().st_size > 0:
            print(f"skip {path.name}")
            continue
        try:
            if suffix == ".docx":
                text = extract_docx(path)
            elif suffix == ".pdf":
                text = extract_pdf(path)
            else:
                print(f"skip unknown {path.name}")
                continue
        except Exception as exc:  # noqa: BLE001
            print(f"FAIL {path.name}: {exc}")
            continue
        out_path.write_text(text, encoding="utf-8", errors="replace")
        print(f"wrote {out_path.name} ({len(text)} chars)")


if __name__ == "__main__":
    main()
